from __future__ import annotations

from collections import Counter
from pathlib import Path
from shutil import copy2
from typing import Iterable

from docxtpl import DocxTemplate
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Alignment, Border, Font, PatternFill, Protection, Side
from copy import copy

from .models import StreetReportContext, TableSheet
from .rules import fmt_datetime, fmt_float, fmt_int


def _normalize_text(text: str) -> str:
    return text.replace("街道街道", "街道").replace("镇街道", "镇").replace("一是. ", "一是").replace("二是. ", "二是").replace("三是. ", "三是")


def _set_paragraph_text_preserve_style(para, text: str) -> None:
    if para.runs:
        para.runs[0].text = text
        for run in para.runs[1:]:
            run.text = ""
    else:
        para.add_run(text)


def _delete_paragraph(para) -> None:
    element = para._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _merge_suggestion_paragraphs(doc) -> bool:
    changed = False
    markers = {"一是", "二是", "三是"}
    paragraphs = list(doc.paragraphs)
    for idx, para in enumerate(paragraphs):
        marker = para.text.strip()
        if marker not in markers:
            continue

        if idx > 0 and not paragraphs[idx - 1].text.strip():
            _delete_paragraph(paragraphs[idx - 1])
            changed = True

        body_idx = idx + 1
        while body_idx < len(paragraphs) and not paragraphs[body_idx].text.strip():
            _delete_paragraph(paragraphs[body_idx])
            changed = True
            body_idx += 1
        if body_idx >= len(paragraphs):
            continue

        body_para = paragraphs[body_idx]
        body_text = body_para.text.strip()
        if not body_text:
            continue

        _set_paragraph_text_preserve_style(body_para, marker + body_text)
        _delete_paragraph(para)
        changed = True

        if body_idx + 1 < len(paragraphs) and not paragraphs[body_idx + 1].text.strip():
            _delete_paragraph(paragraphs[body_idx + 1])
            changed = True
    return changed


def render_docx_template(template_path: Path, output_path: Path, context: dict) -> None:
    doc = DocxTemplate(str(template_path))
    doc.render(context)
    doc.save(str(output_path))
    _postprocess_docx(output_path)


def _postprocess_docx(path: Path) -> None:
    from docx import Document

    doc = Document(path)
    changed = _merge_suggestion_paragraphs(doc)
    for para in doc.paragraphs:
        new_text = _normalize_text(para.text)
        if new_text != para.text:
            _set_paragraph_text_preserve_style(para, new_text)
            changed = True
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    new_text = _normalize_text(para.text)
                    if new_text != para.text:
                        _set_paragraph_text_preserve_style(para, new_text)
                        changed = True
    if changed:
        doc.save(path)


def render_detail_workbook(template_path: Path, output_path: Path, rows: list[dict[str, object]]) -> None:
    wb = load_workbook(template_path)
    ws = wb[wb.sheetnames[0]]
    if not rows:
        wb.save(output_path)
        return

    template_row = 2
    row_style = [copy(ws.cell(template_row, col)._style) for col in range(1, ws.max_column + 1)]
    row_height = ws.row_dimensions[template_row].height

    for i, item in enumerate(rows, start=template_row):
        ws.row_dimensions[i].height = row_height
        for col in range(1, ws.max_column + 1):
            ws.cell(i, col)._style = copy(row_style[col - 1])
        for col, key in enumerate(DETAIL_COLUMN_KEYS, start=1):
            ws.cell(i, col).value = item.get(key, "")

    wb.save(output_path)


DETAIL_COLUMN_KEYS = [
    "编号",
    "案件状态",
    "案件分类",
    "上报时间",
    "截止时间",
    "结案时间",
    "小区",
    "案件描述",
    "案件标签",
    "道路",
    "公园",
    "检查点位1级",
    "检查点位2级",
    "检查点位3级",
    "检查指标1级",
    "检查指标2级",
    "检查指标3级",
    "街镇分中心",
    "区委办局",
    "区委办局二级单位",
    "作业单位",
    "作业单位二级",
    "整改责任单位",
    "处置部门名称一级部门",
    "整改时间二级部门",
]


def render_summary_workbook(output_path: Path, tables: list[TableSheet]) -> None:
    from openpyxl import Workbook

    wb = Workbook()
    if wb.sheetnames:
        wb.remove(wb[wb.sheetnames[0]])
    used_names: set[str] = set()
    for table in tables:
        base = safe_sheet_name(table.title or f"{table.source.stem}-{table.index}")
        name = base
        suffix = 1
        while name in used_names:
            suffix += 1
            name = safe_sheet_name(f"{base[:28]}_{suffix}")
        used_names.add(name)
        ws = wb.create_sheet(title=name)
        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, value in enumerate(row, start=1):
                ws.cell(r_idx, c_idx).value = value
        if table.rows:
            ws.freeze_panes = "A2"
            ws.sheet_view.showGridLines = True
        ws["A1"].font = Font(bold=True)
        _auto_fit_columns(ws)
    wb.save(output_path)


def _auto_fit_columns(ws) -> None:
    max_width = 42
    min_width = 8
    for col_cells in ws.columns:
        column_letter = get_column_letter(col_cells[0].column)
        best = 0
        for cell in col_cells:
            value = cell.value
            if value is None:
                continue
            text = str(value)
            lines = text.splitlines() or [text]
            best = max(best, *(len(line) for line in lines))
        if best == 0:
            continue
        width = min(max(best + 2, min_width), max_width)
        ws.column_dimensions[column_letter].width = width


def safe_sheet_name(name: str) -> str:
    invalid = r'[]:*?/\\'
    for ch in invalid:
        name = name.replace(ch, "_")
    name = name.strip() or "Sheet"
    return name[:31]
